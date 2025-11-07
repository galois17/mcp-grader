import os
import pandas as pd
from docx import Document

# Helper: Excel Reading
def _read_excel_to_text(file_path: str) -> str:
    df = pd.read_excel(file_path, sheet_name=0, header=None)
    lines = []
    for _, row in df.iterrows():
        vals = [str(x) for x in row.tolist() if pd.notna(x)]
        if vals:
            lines.append("\t".join(vals))
    return "\n".join(lines)

def _read_word_to_text(file_path: str) -> str:
    doc = Document(file_path)
    lines = []
    for para in doc.paragraphs:
        txt = para.text.strip()
        if txt:
            lines.append(txt)
    # Include any table cells too
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_text:
                lines.append("\t".join(row_text))
    return "\n".join(lines)

def read_to_text(file_path: str) -> str:
    ext = os.path.splitext(file_path.lower())[1]
    if ext in [".xlsx", ".xls"]:
        return _read_excel_to_text(file_path)
    elif ext == ".docx":
        return _read_word_to_text(file_path)
    elif ext in [".txt"]:
        return open(file_path, "r", encoding="utf-8").read()
    else:
        raise ValueError(f"Unsupported file type: {ext}")

def is_number(value: str) -> bool:
    """Check if a string represents a numeric value."""
    try:
        float(str(value).strip())
        return True
    except Exception:
        return False
    
def numbers_close(a: str, b: str, tol_decimals: int = 2) -> bool:
    """Compare two numeric strings up to specified decimal places."""
    try:
        fa, fb = float(str(a).strip()), float(str(b).strip())
        # Compare after rounding to tolerance
        return round(fa, tol_decimals) == round(fb, tol_decimals)
    except Exception:
        return False

